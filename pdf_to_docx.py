import os
import pathlib
from typing import List

import pypdfium2
from PIL import Image
import pytesseract
import docx


def get_text_from_image(image: Image, lang_set="pol") -> str:
    """
    Returns text detected by TesseractOCR on given image.
    By default, uses Polish language set.
    :param image: Image object with text to detect
    :param lang_set: Language set that will be used by TesseractOCR
    :return: String with text detected on image
    """
    return pytesseract.image_to_string(image, lang=lang_set)


def convert_pdf_to_png(filename: str) -> List[Image.Image]:
    """
    Converts given pdf document to a set of images. Searches for pdf file in "./input".
    :param filename: name of file in directory "./input". Can also be a path starting in "./input"
    :return: list of images made from pages in pdf
    """
    pdf_document = pypdfium2.PdfDocument(f"./input/{filename}")
    number_of_pages = len(pdf_document)
    list_of_images = []
    for page_number in range(number_of_pages):  # converts page in pdf to images
        page = pdf_document.get_page(page_number)
        pil_image = page.render_topil(
            scale=10,
            rotation=0,
            crop=(0, 0, 0, 0),
            greyscale=False,
            optimise_mode=pypdfium2.OptimiseMode.NONE
        )  # render page to PIL Image object
        list_of_images.append(pil_image)
        page.close()
    pdf_document.close()
    return list_of_images


def convert_pdf_to_docx(filename: str) -> str:
    """
    Converts pdf to docx. First, renders pdf as images, then runs them through TesseractOCR to extract text.
    After extraction saves it to docx document with heading indicating what page in pdf it was.
    Every page in pdf is separate page in docx.
    Searches for file in "./input".
    :param filename: Name of file in directory "./input". Can also be a path starting in "./input"
    :return: Name of created docx file
    """
    word_document = docx.Document()
    list_of_images = convert_pdf_to_png(filename)
    for i, image in enumerate(list_of_images):  # iterates over Image list
        extracted_text = get_text_from_image(image)  # extracts text form image
        word_document.add_heading(f"Strona {i}", level=0)  # and adds it on separate pages along with page number
        word_document.add_paragraph(extracted_text)
        word_document.add_page_break()
    # filename is passed with .pdf at the end, so it changes extension to .docx
    filename = filename.replace(".pdf", ".docx")
    word_document.save(f'./output/{filename}')
    return filename


def convert_from_directory(directory: str) -> List[str]:
    """
    Iterates over directory and calls convert_pdf_to_docx for every pdf file in it.
    Searches for directory in "./input"
    :param directory: Name of directory with files to convert. Can also be a path starting in "./input"
    :return: List of files (filenames) created by function, saved in "./output"
    """
    file_list = []
    if not os.path.exists(f"./output/{directory}"):
        os.mkdir(f"./output/{directory}")
    for file in pathlib.Path(f"input/{directory}").iterdir():
        if not file.match("*.pdf"):
            continue
        filepath = str(file).replace(  # convert to str and cut out additional "input\" added by pathlib
            "input\\", "",
            1
        )
        created_file = convert_pdf_to_docx(filepath)
        file_list.append(created_file)
    return file_list


if __name__ == '__main__':
    if not os.path.exists("./input"):
        os.mkdir("./input")
    if not os.path.exists("./output"):
        os.mkdir("./output")
    while True:
        print("PDF to DOCX\n"
              "[1] - konwertuj plik PDF\n"
              "[2] - konwertuj wszystkie pliki w folderze\n"
              "[0] - wyjdz z programu\n"
              "Wybierz rodzaj czynnosci: ")
        action = input()
        match action:
            case "1":
                print("Podaj nazwe pliku znajdujacego sie w folderze input")
                input_file = input()
                if not input_file.endswith(".pdf"):
                    input_file += ".pdf"
                print(f"Konwertuje PDF {input_file} na DOCX")
                result_file = convert_pdf_to_docx(input_file)
                print(f"Zakonczono konwersje, utworzono plik {result_file} w folderze output")
            case "2":
                print("Podaj nazwe folderu z ktorego maja zostac przekonwertowane pliki.\n"
                      "Folder musi znajdowac sie w folderze input.")
                input_directory = input()
                print(f"Konwertuje pliki pdf znajdujace sie w folderze {input_directory}")
                output_file_list = convert_from_directory(input_directory)
                print(f"Zakonczono konwersje, utworzono pliki w folderze output\n"
                      f"Lista plikow:")
                for index, result_filename in enumerate(output_file_list):
                    print(f"{index}\t|\t{result_filename}")
            case "3":
                print("Podaj sciezke do tesseract.exe")
                pytesseract.pytesseract.tesseract_cmd = input()
            case "0":
                break
            case _:
                print("Bledna akcja")
