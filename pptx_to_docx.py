import os
import shutil
import tempfile
import xml.etree.ElementTree as xmltree
import zipfile
from typing import List

from PIL import Image
import pytesseract
import docx
from docx.shared import Cm
import pathlib


def get_text_from_photo(media_path, media_name, word_document, lang_set: str = "pol"):
    """
    Extracts text from image using TesseractOCR. Adds it to docx document along with original photo for reference.
    By default, uses Polish language set.
    :param media_path: path to image
    :param media_name: name of image
    :param word_document: docx.Document object
    :param lang_set: Language set that will be used by TesseractOCR
    :return: None
    """
    photo = Image.open(media_path + media_name)
    text_on_photo = pytesseract.image_to_string(photo, lang=lang_set)
    word_document.add_paragraph(text_on_photo)  # extracted text
    word_document.add_heading(f"Zródło tekstu - {media_name}", level=1)  # source of text
    word_document.add_picture(media_path + media_name, width=Cm(15))


def find_images_in_slide(slide_path, relation_path, media_path, word_document):
    """
    Searches for images in slides and extracts text form them.
    :param slide_path: path to xml file containing slide
    :param relation_path: path to file with relations in slide
    :param media_path: path to media files in pptx presentation
    :param word_document: docx.Document object
    :return: None
    """
    file = xmltree.parse(slide_path)
    relations = xmltree.parse(relation_path)
    file_root = file.getroot()
    rel_root = relations.getroot()
    for blip in file_root.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
        photo_id = blip.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed']
        for child in rel_root:
            if photo_id == child.attrib["Id"]:
                media_name = child.attrib["Target"].split("/")[2]
                if media_name.endswith((".png", ".jpg", ".jpeg")):
                    get_text_from_photo(media_path, media_name, word_document)


def find_text_in_slide(slide_path, word_document):
    """
    Searches slide xml file for text and adds it to word document using function parse_text()
    :param slide_path: path to xml file containing slide
    :param word_document: docx.Document object
    :return: None
    """
    file = xmltree.parse(slide_path)
    root = file.getroot()
    paragraph_iter = root.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}p")
    for p in paragraph_iter:
        paragraph = word_document.add_paragraph()
        for r in p.findall("{http://schemas.openxmlformats.org/drawingml/2006/main}r"):
            parse_text(r, paragraph)


def parse_text(text_data, paragraph):
    """
    Extracts text from slide and adds it to word document. Preserves bold text in slide
    :param text_data: xml element containing text on slide
    :param paragraph: paragraph of docx.Document object that will get text added
    :return: None
    """
    rPr = text_data.find("{http://schemas.openxmlformats.org/drawingml/2006/main}rPr")
    t = text_data.find("{http://schemas.openxmlformats.org/drawingml/2006/main}t")
    if "b" in rPr.attrib.keys():
        if rPr.attrib["b"]:
            if t.text is not None:
                paragraph.add_run(t.text.strip() + " ").bold = True
    else:
        if t.text is not None:
            paragraph.add_run(t.text.strip() + " ")


def count_slides(slides_path: str) -> int:
    """
    Counts slides in pptx presentation.
    :param slides_path: path to folder containing slides xml flies
    :return: Number of slides in pptx presentation
    """
    count = 0
    for path in pathlib.Path(slides_path).iterdir():
        if path.is_file() and str(path).endswith(".xml"):
            count += 1
    return count


# noinspection PyTypeChecker
def unzip_pptx(file_name: str, temp_dir: tempfile.TemporaryDirectory, file_path: str = ""):
    """
    Copies pptx file to temporary directory and changes it to zip file, then extracts files from zip in temporary directory.
    :param file_name: name of the file to extract
    :param temp_dir: path to temporary directory
    :param file_path: path to file (if in directory inside "./input")
    :return: None
    """
    temp_file_path = shutil.copy(f"./input/{file_path}{file_name}", f"{temp_dir}/{file_name.replace('.pptx', '.zip')}")
    zipfile.ZipFile(temp_file_path).extractall(temp_dir)


def convert_pptx_to_docx(file_name: str, file_path: str = "") -> str:
    """
    Runs all neccesary functions to convert pptx presentation to docx document.
    Extracts files from pptx, then iterates over all slides and copies text form slide and images into docx document.
    :param file_name: name of the file to extract
    :param file_path: path to file (if in directory inside "./input")
    :return: name of created docx document in directory "./output"
    """
    word_document = docx.Document()
    with tempfile.TemporaryDirectory() as temp_dir:
        unzip_pptx(f"{file_name}", temp_dir, file_path=file_path)
        # as a precaution, the program checks for important folders and creates then if they don't exist
        # they can be empty - if they are, there's just nothing to extract and the program won't do anything
        if not os.path.exists(f"{temp_dir}\\ppt\\slides\\"):
            os.mkdir(f"{temp_dir}\\ppt\\slides\\")
        slides_path = f"{temp_dir}\\ppt\\slides/"
        if not os.path.exists(f"{temp_dir}\\ppt\\media\\"):
            os.mkdir(f"{temp_dir}\\ppt\\media\\")
        media_path = f"{temp_dir}\\ppt\\media\\"
        if not os.path.exists(f"{temp_dir}\\ppt\\slides\\_rels\\"):
            os.mkdir(f"{temp_dir}/ppt\\slides\\_rels\\")
        relations_path = f"{temp_dir}\\ppt\\slides\\_rels\\"
        for i in range(1, count_slides(slides_path) + 1):  # iterate over all slides in presentation
            word_document.add_heading(f"Slajd {i} - tekst", level=0)
            find_text_in_slide(
                slide_path=f"{slides_path}slide{i}.xml",
                word_document=word_document
            )
            word_document.add_page_break()
            word_document.add_heading(f"Slajd {i} - zdjęcia", level=0)
            find_images_in_slide(
                slide_path=f"{slides_path}slide{i}.xml",
                relation_path=f"{relations_path}slide{i}.xml.rels",
                media_path=media_path,
                word_document=word_document
            )
            word_document.add_page_break()
    file_name = file_name.replace(".pptx", ".docx")
    word_document.save(f'./output/{file_path}{file_name}')
    return file_name


def convert_from_directory(directory: str) -> List[str]:
    """
    Iterates over directory and calls convert_pptx_to_docx for every pdf file in it.
    Searches for directory in "./input"
    :param directory: Name of directory with files to convert. Can also be a path starting in "./input"
    :return: List of files (filenames) created by function, saved in "./output"
    """
    file_list = []
    if not os.path.exists(f"./output/{directory}"):
        os.mkdir(f"./output/{directory}")
    for file in pathlib.Path(f"input/{directory}").iterdir():
        if not file.match("*.pptx"):
            continue
        # convert to str, cut out additional "input\" added by pathlib and remove filename from path
        file_path = str(file).replace("input\\", "", 1).replace(file.name, "")
        if not file_path.endswith("/"):
            file_path += "/"
        created_file = convert_pptx_to_docx(file.name, file_path=file_path)
        file_list.append(created_file)
    return file_list


if __name__ == '__main__':
    if not os.path.exists("./input"):
        os.mkdir("./input")
    if not os.path.exists("./output"):
        os.mkdir("./output")
    while True:
        print("PPTX to DOCX\n"
              "[1] - konwertuj plik PPTX\n"
              "[2] - konwertuj wszystkie pliki w folderze\n"
              "[3] - ustaw sciezke do TesseractOCR"
              "[0] - wyjdz z programu\n"
              "Wybierz rodzaj czynnosci: ")
        action = input()
        match action:
            case "1":
                print("Podaj nazwe pliku znajdujacego sie w folderze input")
                input_file = input()
                if not input_file.endswith(".pptx"):
                    input_file += ".pptx"
                print(f"Konwertuje PPTX {input_file} na DOCX")
                result_file = convert_pptx_to_docx(input_file)
                print(f"Zakonczono konwersje, utworzono plik {result_file} w folderze output")
            case "2":
                print("Podaj nazwe folderu z ktorego maja zostac przekonwertowane pliki.\n"
                      "Folder musi znajdowac sie w folderze input.")
                input_directory = input()
                print(f"Konwertuje pliki PPTX znajdujace sie w folderze {input_directory}")
                output_file_list = convert_from_directory(input_directory)
                print(f"Zakonczono konwersje, utworzono pliki w folderze output\n"
                      f"Lista plikow:")
                for index, result_filename in enumerate(output_file_list):
                    print(f"{index}\t|\t{result_filename}")
            case "3":
                print("Podaj sciezke do tesseract.exe")
                pytesseract.pytesseract.tesseract_cmd = input()
            case "0":
                break
            case _:
                print("Bledna akcja")
